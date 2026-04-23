# import streamlit as st
# from gtts import gTTS
# from deep_translator import GoogleTranslator
# import os
# import PyPDF2
# import docx
# import time

# # 🎨 Page Config
# st.set_page_config(page_title="AI Voice Assistant", page_icon="🎤", layout="centered")

# # 🎨 Custom UI Styling
# st.markdown("""
#     <style>
#     body {
#         background-color: #0f172a;
#         color: white;
#     }
#     .stApp {
#         background: linear-gradient(135deg, #1e3a8a, #9333ea);
#         color: white;
#     }
#     .stButton>button {
#         background-color: #22c55e;
#         color: white;
#         border-radius: 10px;
#         height: 3em;
#         width: 100%;
#         font-size: 18px;
#     }
#     .stTextArea textarea {
#         background-color: #020617;
#         color: white;
#     }
#     </style>
# """, unsafe_allow_html=True)

# # 🎤 Title
# st.title("🎤 Smart AI Voice Assistant")
# st.write("✨ Convert text & files into natural speech")

# # 📄 File Reader
# def read_file(file):
#     try:
#         if file.name.endswith(".txt"):
#             return file.read().decode("utf-8")

#         elif file.name.endswith(".pdf"):
#             reader = PyPDF2.PdfReader(file)
#             text = ""
#             for page in reader.pages:
#                 if page.extract_text():
#                     text += page.extract_text() + "\n"
#             return text

#         elif file.name.endswith(".docx"):
#             doc = docx.Document(file)
#             return "\n".join([p.text for p in doc.paragraphs])

#         else:
#             return ""

#     except Exception as e:
#         st.error(f"❌ File Error: {e}")
#         return ""

# # 🌍 Language selection
# lang = st.selectbox(
#     "🌍 Select Language",
#     ["English", "Hindi", "Hinglish", "Bhojpuri"]
# )

# # 📂 Upload file
# uploaded_file = st.file_uploader("📂 Upload file", type=["txt", "pdf", "docx"])

# text = ""

# if uploaded_file:
#     text = read_file(uploaded_file)
#     st.success("✅ File Loaded Successfully!")

# # 📝 Text input
# text = st.text_area("📝 Enter Text", value=text, height=200)

# # 🔄 Translation
# def convert_text(text, lang):
#     try:
#         if lang == "Hindi":
#             return GoogleTranslator(source='auto', target='hi').translate(text)

#         elif lang == "Hinglish":
#             return GoogleTranslator(source='auto', target='hi').translate(text)

#         elif lang == "Bhojpuri":
#             return GoogleTranslator(source='auto', target='hi').translate(text)

#         return text

#     except Exception as e:
#         st.error(f"❌ Translation Error: {e}")
#         return text

# # 🎚️ Extra feature: Speed control
# speed = st.slider("⚡ Voice Speed", 0.5, 2.0, 1.0)

# # 🎵 Generate Voice
# if st.button("🎧 Convert to Voice"):
#     if text.strip():
#         final_text = convert_text(text, lang)

#         try:
#             lang_code = "hi" if lang != "English" else "en"

#             with st.spinner("⏳ Generating voice..."):
#                 tts = gTTS(text=final_text, lang=lang_code, slow=(speed < 1))
#                 tts.save("output.mp3")
#                 time.sleep(1)

#             st.success("✅ Voice Generated Successfully!")

#             # 🔊 Audio Player
#             st.audio("output.mp3")

#         except Exception as e:
#             st.error(f"❌ TTS Error: {e}")

#     else:
#         st.warning("⚠️ Please enter some text!")

# # 💾 Download
# if os.path.exists("output.mp3"):
#     with open("output.mp3", "rb") as f:
#         st.download_button(
#             "⬇️ Download Audio",
#             f,
#             file_name="voice.mp3"
#         )

# # 🧹 Clear button
# if st.button("🗑 Clear"):
#     if os.path.exists("output.mp3"):
#         os.remove("output.mp3")
#     st.rerun()

# # Footer
# st.markdown("---")
# st.caption("🚀 Made By me 😎")
import streamlit as st
from gtts import gTTS
from deep_translator import GoogleTranslator
import os
import PyPDF2
import docx
import time

# ---------------- CONFIG ----------------
st.set_page_config(page_title="AI Voice Assistant", page_icon="🎤")

# ---------------- STYLE ----------------
st.markdown("""
<style>
.stApp {
    background: linear-gradient(135deg, #1e3a8a, #9333ea);
    color: white;
}
.stButton>button {
    background-color: #22c55e;
    color: white;
    border-radius: 8px;
}
</style>
""", unsafe_allow_html=True)

st.title("🎤 Smart AI Voice Assistant")

# ---------------- FILE READER ----------------
def read_file(file):
    if file.name.endswith(".txt"):
        return file.read().decode("utf-8")

    elif file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            if page.extract_text():
                text += page.extract_text()
        return text

    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])

    return ""

# ---------------- TRANSLATE ----------------
def translate_text(text, lang):
    try:
        if lang == "Hindi":
            return GoogleTranslator(source='auto', target='hi').translate(text)
        return text
    except:
        return text

# ---------------- UI ----------------
lang = st.selectbox("🌍 Select Language", ["English", "Hindi"])

uploaded_file = st.file_uploader("📂 Upload File", ["txt", "pdf", "docx"])

text = ""

if uploaded_file:
    text = read_file(uploaded_file)
    st.success("✅ File Loaded")

# Text input
text = st.text_area("📝 Enter Text", value=text, height=200)

# Stats
if text:
    words = len(text.split())
    st.info(f"📊 Words: {words} | ⏱ Time: {words//2} sec")

# Convert button
if st.button("🎧 Convert to Voice"):
    if text.strip():
        final_text = translate_text(text, lang)

        try:
            # create audio folder if not exists
            os.makedirs("audio", exist_ok=True)

            filename = f"audio_{int(time.time())}.mp3"
            filepath = os.path.join("audio", filename)

            tts = gTTS(text=final_text, lang="hi" if lang=="Hindi" else "en")
            tts.save(filepath)

            st.success("✅ Voice Generated")
            st.audio(filepath)

            # download
            with open(filepath, "rb") as f:
                st.download_button("⬇️ Download", f, file_name=filename)

        except Exception as e:
            st.error(f"❌ Error: {e}")
    else:
        st.warning("⚠️ Enter text first")

# ---------------- HISTORY ----------------
st.subheader("📁 Recent Audio")

if os.path.exists("audio"):
    files = sorted(os.listdir("audio"), reverse=True)

    for file in files[:5]:
        st.audio(os.path.join("audio", file))